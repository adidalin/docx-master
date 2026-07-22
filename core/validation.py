"""
验证模块
提供文档结构验证、格式验证和自动修复功能
"""

from typing import Dict, List, Optional, Any
from docx import Document


class DocumentValidator:
    """文档验证器"""
    
    def __init__(self, document: Document):
        """
        初始化验证器
        
        Args:
            document: Document对象
        """
        self.document = document
        self.errors = []
        self.warnings = []
    
    def validate_structure(self) -> Dict[str, Any]:
        """
        验证文档结构
        
        Returns:
            验证结果
        """
        result = {
            "valid": True,
            "errors": [],
            "warnings": [],
            "statistics": {}
        }
        
        # 检查段落
        paragraphs = self.document.paragraphs
        result["statistics"]["paragraph_count"] = len(paragraphs)
        
        # 检查表格
        tables = self.document.tables
        result["statistics"]["table_count"] = len(tables)
        
        # 检查章节
        sections = self.document.sections
        result["statistics"]["section_count"] = len(sections)
        
        # 检查样式
        styles = self.document.styles
        result["statistics"]["style_count"] = len(styles)
        
        # 检查标题层级
        heading_levels = set()
        for para in paragraphs:
            if para.style.name.startswith("Heading"):
                try:
                    level = int(para.style.name.replace("Heading", ""))
                    heading_levels.add(level)
                except ValueError:
                    pass
        
        result["statistics"]["heading_levels"] = sorted(list(heading_levels))
        
        # 检查空段落
        empty_paragraphs = [i for i, p in enumerate(paragraphs) if not p.text.strip()]
        if empty_paragraphs:
            result["warnings"].append({
                "type": "empty_paragraphs",
                "message": f"发现{len(empty_paragraphs)}个空段落",
                "locations": empty_paragraphs
            })
        
        # 检查标题顺序
        if heading_levels:
            max_level = max(heading_levels)
            min_level = min(heading_levels)
            if max_level - min_level + 1 != len(heading_levels):
                result["warnings"].append({
                    "type": "heading_level_gap",
                    "message": "标题层级不连续",
                    "levels": heading_levels
                })
        
        return result
    
    def validate_formatting(self) -> Dict[str, Any]:
        """
        验证格式
        
        Returns:
            验证结果
        """
        result = {
            "valid": True,
            "errors": [],
            "warnings": []
        }
        
        # 检查字体一致性
        font_issues = self._check_font_consistency()
        if font_issues:
            result["warnings"].extend(font_issues)
        
        # 检查段落格式一致性
        paragraph_issues = self._check_paragraph_consistency()
        if paragraph_issues:
            result["warnings"].extend(paragraph_issues)
        
        # 检查页面格式
        page_issues = self._check_page_format()
        if page_issues:
            result["warnings"].extend(page_issues)
        
        return result
    
    def _check_font_consistency(self) -> List[Dict[str, Any]]:
        """
        检查字体一致性
        
        Returns:
            字体问题列表
        """
        issues = []
        fonts = set()
        
        for para in self.document.paragraphs:
            for run in para.runs:
                if run.font.name:
                    fonts.add(run.font.name)
        
        if len(fonts) > 5:
            issues.append({
                "type": "too_many_fonts",
                "message": f"使用了{len(fonts)}种不同的字体",
                "fonts": list(fonts)
            })
        
        return issues
    
    def _check_paragraph_consistency(self) -> List[Dict[str, Any]]:
        """
        检查段落格式一致性
        
        Returns:
            段落问题列表
        """
        issues = []
        alignments = set()
        
        for para in self.document.paragraphs:
            if para.alignment:
                alignments.add(str(para.alignment))
        
        if len(alignments) > 3:
            issues.append({
                "type": "too_many_alignments",
                "message": f"使用了{len(alignments)}种不同的对齐方式",
                "alignments": list(alignments)
            })
        
        return issues
    
    def _check_page_format(self) -> List[Dict[str, Any]]:
        """
        检查页面格式
        
        Returns:
            页面问题列表
        """
        issues = []
        
        for i, section in enumerate(self.document.sections):
            # 检查页边距
            if section.left_margin and section.right_margin:
                if section.left_margin == section.right_margin:
                    pass  # 正常
                # 可以添加更多检查
            
            # 检查纸张大小
            if section.page_width and section.page_height:
                # A4: 210mm x 297mm
                if abs(section.page_width - 210 * 3600) > 1000 or abs(section.page_height - 297 * 3600) > 1000:
                    issues.append({
                        "type": "non_standard_page_size",
                        "message": f"第{i+1}节使用了非标准纸张大小",
                        "width": section.page_width,
                        "height": section.page_height
                    })
        
        return issues
    
    def validate_for_official_document(self) -> Dict[str, Any]:
        """
        验证公文格式
        
        Returns:
            验证结果
        """
        result = {
            "valid": True,
            "errors": [],
            "warnings": []
        }
        
        # 检查页面设置
        if self.document.sections:
            section = self.document.sections[0]
            
            # 检查纸张大小
            if section.page_width and section.page_height:
                # A4: 210mm x 297mm
                expected_width = 210 * 3600
                expected_height = 297 * 3600
                if abs(section.page_width - expected_width) > 1000:
                    result["warnings"].append({
                        "type": "incorrect_page_width",
                        "message": "纸张宽度不符合A4标准",
                        "expected": "210mm",
                        "actual": f"{section.page_width / 3600:.1f}mm"
                    })
            
            # 检查页边距
            if section.left_margin:
                expected_left = 2.8 * 3600
                if abs(section.left_margin - expected_left) > 1000:
                    result["warnings"].append({
                        "type": "incorrect_left_margin",
                        "message": "左边距不符合公文标准",
                        "expected": "2.8cm",
                        "actual": f"{section.left_margin / 3600:.1f}cm"
                    })
        
        return result
    
    def auto_fix(self) -> Dict[str, Any]:
        """
        自动修复
        
        Returns:
            修复结果
        """
        result = {
            "fixed": [],
            "skipped": [],
            "failed": []
        }
        
        # 修复空段落
        empty_indices = [i for i, p in enumerate(self.document.paragraphs) if not p.text.strip()]
        for i in reversed(empty_indices):
            try:
                # 删除空段落
                p = self.document.paragraphs[i]
                p._element.getparent().remove(p._element)
                result["fixed"].append(f"删除空段落 {i}")
            except Exception as e:
                result["failed"].append(f"删除空段落 {i} 失败: {str(e)}")
        
        return result


def validate_document(doc_path: str) -> Dict[str, Any]:
    """
    验证文档
    
    Args:
        doc_path: 文档路径
        
    Returns:
        验证结果
    """
    try:
        doc = Document(doc_path)
        validator = DocumentValidator(doc)
        
        structure_result = validator.validate_structure()
        formatting_result = validator.validate_formatting()
        
        return {
            "structure": structure_result,
            "formatting": formatting_result,
            "valid": structure_result["valid"] and formatting_result["valid"]
        }
    except Exception as e:
        return {
            "valid": False,
            "error": str(e)
        }


def validate_official_document(doc_path: str) -> Dict[str, Any]:
    """
    验证公文格式
    
    Args:
        doc_path: 文档路径
        
    Returns:
        验证结果
    """
    try:
        doc = Document(doc_path)
        validator = DocumentValidator(doc)
        
        result = validator.validate_for_official_document()
        return result
    except Exception as e:
        return {
            "valid": False,
            "error": str(e)
        }
