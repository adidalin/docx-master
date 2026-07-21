"""
样式管理模块
提供样式的创建、应用、继承等功能
"""

from typing import Dict, List, Optional, Any
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class StyleManager:
    """样式管理器，提供文档样式管理功能"""
    
    def __init__(self, document: Document):
        """
        初始化样式管理器
        
        Args:
            document: Document对象
        """
        self.document = document
        
    def get_style(self, style_name: str) -> Optional[Any]:
        """
        获取样式
        
        Args:
            style_name: 样式名称
            
        Returns:
            样式对象
        """
        try:
            return self.document.styles[style_name]
        except KeyError:
            return None
    
    def get_all_styles(self) -> List[Dict[str, Any]]:
        """
        获取所有样式
        
        Returns:
            样式信息列表
        """
        styles = []
        for style in self.document.styles:
            style_info = {
                "name": style.name,
                "type": str(style.type),
                "builtin": style.builtin,
                "hidden": style.hidden,
                "priority": style.priority
            }
            styles.append(style_info)
            
        return styles
    
    def create_paragraph_style(self, style_name: str, 
                              font_name: Optional[str] = None,
                              font_size: Optional[int] = None,
                              bold: Optional[bool] = None,
                              italic: Optional[bool] = None,
                              color: Optional[str] = None,
                              alignment: Optional[str] = None,
                              line_spacing: Optional[float] = None,
                              space_before: Optional[int] = None,
                              space_after: Optional[int] = None,
                              first_line_indent: Optional[int] = None,
                              left_indent: Optional[int] = None,
                              right_indent: Optional[int] = None,
                              base_style: Optional[str] = None,
                              next_style: Optional[str] = None) -> Any:
        """
        创建段落样式
        
        Args:
            style_name: 样式名称
            font_name: 字体名称
            font_size: 字体大小
            bold: 是否加粗
            italic: 是否斜体
            color: 颜色
            alignment: 对齐方式
            line_spacing: 行距
            space_before: 段前间距
            space_after: 段后间距
            first_line_indent: 首行缩进
            left_indent: 左缩进
            right_indent: 右缩进
            base_style: 基础样式
            next_style: 下一段落样式
            
        Returns:
            样式对象
        """
        # 检查样式是否已存在
        if style_name in [s.name for s in self.document.styles]:
            print(f"样式已存在: {style_name}")
            return self.document.styles[style_name]
            
        # 创建样式
        style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        
        # 设置基础样式
        if base_style:
            try:
                style.base_style = self.document.styles[base_style]
            except KeyError:
                print(f"基础样式不存在: {base_style}")
                
        # 设置下一段落样式
        if next_style:
            try:
                style.next_style = self.document.styles[next_style]
            except KeyError:
                print(f"下一段落样式不存在: {next_style}")
                
        # 设置字体格式
        if font_name:
            style.font.name = font_name
            style.font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            
        if font_size:
            style.font.size = Pt(font_size)
            
        if bold is not None:
            style.font.bold = bold
            
        if italic is not None:
            style.font.italic = italic
            
        if color:
            color = color.lstrip('#')
            style.font.color.rgb = RGBColor.from_string(color)
            
        # 设置段落格式
        if alignment:
            alignment_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            if alignment.lower() in alignment_map:
                style.paragraph_format.alignment = alignment_map[alignment.lower()]
                
        if line_spacing is not None:
            style.paragraph_format.line_spacing = line_spacing
            
        if space_before is not None:
            style.paragraph_format.space_before = Pt(space_before)
            
        if space_after is not None:
            style.paragraph_format.space_after = Pt(space_after)
            
        if first_line_indent is not None:
            style.paragraph_format.first_line_indent = Pt(first_line_indent)
            
        if left_indent is not None:
            style.paragraph_format.left_indent = Pt(left_indent)
            
        if right_indent is not None:
            style.paragraph_format.right_indent = Pt(right_indent)
            
        return style
    
    def create_character_style(self, style_name: str,
                              font_name: Optional[str] = None,
                              font_size: Optional[int] = None,
                              bold: Optional[bool] = None,
                              italic: Optional[bool] = None,
                              color: Optional[str] = None,
                              underline: Optional[bool] = None,
                              strikethrough: Optional[bool] = None,
                              superscript: Optional[bool] = None,
                              subscript: Optional[bool] = None,
                              base_style: Optional[str] = None) -> Any:
        """
        创建字符样式
        
        Args:
            style_name: 样式名称
            font_name: 字体名称
            font_size: 字体大小
            bold: 是否加粗
            italic: 是否斜体
            color: 颜色
            underline: 是否下划线
            strikethrough: 是否删除线
            superscript: 是否上标
            subscript: 是否下标
            base_style: 基础样式
            
        Returns:
            样式对象
        """
        # 检查样式是否已存在
        if style_name in [s.name for s in self.document.styles]:
            print(f"样式已存在: {style_name}")
            return self.document.styles[style_name]
            
        # 创建样式
        style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
        
        # 设置基础样式
        if base_style:
            try:
                style.base_style = self.document.styles[base_style]
            except KeyError:
                print(f"基础样式不存在: {base_style}")
                
        # 设置字体格式
        if font_name:
            style.font.name = font_name
            style.font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            
        if font_size:
            style.font.size = Pt(font_size)
            
        if bold is not None:
            style.font.bold = bold
            
        if italic is not None:
            style.font.italic = italic
            
        if color:
            color = color.lstrip('#')
            style.font.color.rgb = RGBColor.from_string(color)
            
        if underline is not None:
            style.font.underline = underline
            
        if strikethrough is not None:
            style.font.strikethrough = strikethrough
            
        if superscript is not None:
            style.font.superscript = superscript
            
        if subscript is not None:
            style.font.subscript = subscript
            
        return style
    
    def create_table_style(self, style_name: str,
                          alignment: Optional[str] = None,
                          borders: Optional[bool] = None,
                          border_color: Optional[str] = None,
                          border_size: Optional[int] = None,
                          header_row: Optional[bool] = None,
                          banded_rows: Optional[bool] = None,
                          banded_columns: Optional[bool] = None,
                          base_style: Optional[str] = None) -> Any:
        """
        创建表格样式
        
        Args:
            style_name: 样式名称
            alignment: 表格对齐方式
            borders: 是否显示边框
            border_color: 边框颜色
            border_size: 边框大小
            header_row: 是否有标题行
            banded_rows: 是否交替行
            banded_columns: 是否交替列
            base_style: 基础样式
            
        Returns:
            样式对象
        """
        # 检查样式是否已存在
        if style_name in [s.name for s in self.document.styles]:
            print(f"样式已存在: {style_name}")
            return self.document.styles[style_name]
            
        # 创建样式
        style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.TABLE)
        
        # 设置基础样式
        if base_style:
            try:
                style.base_style = self.document.styles[base_style]
            except KeyError:
                print(f"基础样式不存在: {base_style}")
                
        # 设置表格格式
        if alignment:
            alignment_map = {
                "left": WD_TABLE_ALIGNMENT.LEFT,
                "center": WD_TABLE_ALIGNMENT.CENTER,
                "right": WD_TABLE_ALIGNMENT.RIGHT
            }
            if alignment.lower() in alignment_map:
                style.alignment = alignment_map[alignment.lower()]
                
        # 设置边框
        if borders is not None:
            tbl = style._element
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            
            if borders:
                borders_element = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    
                    if border_color:
                        border.set(qn('w:color'), border_color.lstrip('#'))
                        
                    if border_size:
                        border.set(qn('w:sz'), str(border_size * 2))
                        
                    border.set(qn('w:space'), '0')
                    borders_element.append(border)
                    
                tblPr.append(borders_element)
            else:
                borders_element = tblPr.find(qn('w:tblBorders'))
                if borders_element is not None:
                    tblPr.remove(borders_element)
                    
        # 设置标题行
        if header_row is not None:
            tbl = style._element
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            
            look = OxmlElement('w:tblStyleRowBandSize')
            look.set(qn('w:val'), '1' if header_row else '0')
            tblPr.append(look)
            
        # 设置交替行
        if banded_rows is not None:
            tbl = style._element
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            
            look = OxmlElement('w:tblStyleRowBandSize')
            look.set(qn('w:val'), '1' if banded_rows else '0')
            tblPr.append(look)
            
        # 设置交替列
        if banded_columns is not None:
            tbl = style._element
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            
            look = OxmlElement('w:tblStyleColBandSize')
            look.set(qn('w:val'), '1' if banded_columns else '0')
            tblPr.append(look)
            
        return style
    
    def apply_style(self, paragraph, style_name: str):
        """
        应用样式
        
        Args:
            paragraph: Paragraph对象
            style_name: 样式名称
        """
        try:
            paragraph.style = self.document.styles[style_name]
        except KeyError:
            print(f"样式不存在: {style_name}")
    
    def modify_style(self, style_name: str, **kwargs):
        """
        修改样式
        
        Args:
            style_name: 样式名称
            **kwargs: 要修改的属性
        """
        style = self.get_style(style_name)
        if not style:
            print(f"样式不存在: {style_name}")
            return
            
        # 修改字体属性
        if 'font_name' in kwargs:
            style.font.name = kwargs['font_name']
            style.font.element.rPr.rFonts.set(qn('w:eastAsia'), kwargs['font_name'])
            
        if 'font_size' in kwargs:
            style.font.size = Pt(kwargs['font_size'])
            
        if 'bold' in kwargs:
            style.font.bold = kwargs['bold']
            
        if 'italic' in kwargs:
            style.font.italic = kwargs['italic']
            
        if 'color' in kwargs:
            color = kwargs['color'].lstrip('#')
            style.font.color.rgb = RGBColor.from_string(color)
            
        # 修改段落属性
        if 'alignment' in kwargs:
            alignment_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            if kwargs['alignment'].lower() in alignment_map:
                style.paragraph_format.alignment = alignment_map[kwargs['alignment'].lower()]
                
        if 'line_spacing' in kwargs:
            style.paragraph_format.line_spacing = kwargs['line_spacing']
            
        if 'space_before' in kwargs:
            style.paragraph_format.space_before = Pt(kwargs['space_before'])
            
        if 'space_after' in kwargs:
            style.paragraph_format.space_after = Pt(kwargs['space_after'])
            
        if 'first_line_indent' in kwargs:
            style.paragraph_format.first_line_indent = Pt(kwargs['first_line_indent'])
            
        if 'left_indent' in kwargs:
            style.paragraph_format.left_indent = Pt(kwargs['left_indent'])
            
        if 'right_indent' in kwargs:
            style.paragraph_format.right_indent = Pt(kwargs['right_indent'])
    
    def delete_style(self, style_name: str):
        """
        删除样式
        
        Args:
            style_name: 样式名称
        """
        style = self.get_style(style_name)
        if not style:
            print(f"样式不存在: {style_name}")
            return
            
        # 检查是否是内置样式
        if style.builtin:
            print(f"不能删除内置样式: {style_name}")
            return
            
        # 检查是否有段落使用此样式
        for para in self.document.paragraphs:
            if para.style.name == style_name:
                print(f"样式正在被使用，不能删除: {style_name}")
                return
                
        # 删除样式
        self.document.styles.remove(style)
    
    def import_styles(self, source_doc_path: str, style_names: Optional[List[str]] = None):
        """
        从另一个文档导入样式
        
        Args:
            source_doc_path: 源文档路径
            style_names: 要导入的样式名称列表，如果不提供则导入所有样式
        """
        try:
            source_doc = Document(source_doc_path)
            
            for style in source_doc.styles:
                if style_names and style.name not in style_names:
                    continue
                    
                # 检查样式是否已存在
                if style.name in [s.name for s in self.document.styles]:
                    continue
                    
                # 复制样式
                new_style = self.document.styles.add_style(style.name, style.type)
                
                # 复制属性
                if hasattr(style, 'font') and style.font:
                    new_style.font.name = style.font.name
                    new_style.font.size = style.font.size
                    new_style.font.bold = style.font.bold
                    new_style.font.italic = style.font.italic
                    
                if hasattr(style, 'paragraph_format') and style.paragraph_format:
                    new_style.paragraph_format.alignment = style.paragraph_format.alignment
                    new_style.paragraph_format.line_spacing = style.paragraph_format.line_spacing
                    new_style.paragraph_format.space_before = style.paragraph_format.space_before
                    new_style.paragraph_format.space_after = style.paragraph_format.space_after
                    
        except Exception as e:
            print(f"导入样式失败: {e}")
    
    def export_styles(self, output_path: str, style_names: Optional[List[str]] = None):
        """
        导出样式到新文档
        
        Args:
            output_path: 输出路径
            style_names: 要导出的样式名称列表，如果不提供则导出所有样式
        """
        try:
            # 创建新文档
            new_doc = Document()
            
            # 复制样式
            for style in self.document.styles:
                if style_names and style.name not in style_names:
                    continue
                    
                # 检查样式是否已存在
                if style.name in [s.name for s in new_doc.styles]:
                    continue
                    
                # 复制样式
                new_style = new_doc.styles.add_style(style.name, style.type)
                
                # 复制属性
                if hasattr(style, 'font') and style.font:
                    new_style.font.name = style.font.name
                    new_style.font.size = style.font.size
                    new_style.font.bold = style.font.bold
                    new_style.font.italic = style.font.italic
                    
                if hasattr(style, 'paragraph_format') and style.paragraph_format:
                    new_style.paragraph_format.alignment = style.paragraph_format.alignment
                    new_style.paragraph_format.line_spacing = style.paragraph_format.line_spacing
                    new_style.paragraph_format.space_before = style.paragraph_format.space_before
                    new_style.paragraph_format.space_after = style.paragraph_format.space_after
                    
            # 保存文档
            new_doc.save(output_path)
            
        except Exception as e:
            print(f"导出样式失败: {e}")
    
    def get_style_usage(self, style_name: str) -> List[int]:
        """
        获取样式的使用情况
        
        Args:
            style_name: 样式名称
            
        Returns:
            使用该样式的段落索引列表
        """
        usage = []
        for i, para in enumerate(self.document.paragraphs):
            if para.style.name == style_name:
                usage.append(i)
                
        return usage
    
    def rename_style(self, old_name: str, new_name: str):
        """
        重命名样式
        
        Args:
            old_name: 原样式名称
            new_name: 新样式名称
        """
        style = self.get_style(old_name)
        if not style:
            print(f"样式不存在: {old_name}")
            return
            
        # 检查新名称是否已存在
        if new_name in [s.name for s in self.document.styles]:
            print(f"样式名称已存在: {new_name}")
            return
            
        # 重命名样式
        style.name = new_name
        
        # 更新所有使用该样式的段落
        for para in self.document.paragraphs:
            if para.style.name == old_name:
                para.style = self.document.styles[new_name]