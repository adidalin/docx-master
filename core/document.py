"""
文档管理模块
提供文档的创建、打开、保存、分析等核心功能
"""

import os
import json
from typing import Dict, List, Optional, Union, Any
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .formatting import FormattingManager
from .styles import StyleManager
from .ai_integration import AIIntegration
from .utils import Utils


class DocumentManager:
    """文档管理器，提供文档的核心操作功能"""
    
    def __init__(self, template_path: Optional[str] = None):
        """
        初始化文档管理器
        
        Args:
            template_path: 模板文件路径，如果不提供则创建空白文档
        """
        self.template_path = template_path
        self.document = None
        self.formatting_manager = None
        self.style_manager = None
        self.ai_integration = None
        self.utils = Utils()
        
    def create_document(self, template_path: Optional[str] = None) -> Document:
        """
        创建新文档
        
        Args:
            template_path: 模板文件路径
            
        Returns:
            Document对象
        """
        if template_path and os.path.exists(template_path):
            self.document = Document(template_path)
        elif self.template_path and os.path.exists(self.template_path):
            self.document = Document(self.template_path)
        else:
            self.document = Document()
            
        # 初始化管理器
        self.formatting_manager = FormattingManager(self.document)
        self.style_manager = StyleManager(self.document)
        self.ai_integration = AIIntegration()
        
        return self.document
    
    def open_document(self, file_path: str) -> Document:
        """
        打开现有文档
        
        Args:
            file_path: 文档路径
            
        Returns:
            Document对象
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文档不存在: {file_path}")
            
        self.document = Document(file_path)
        
        # 初始化管理器
        self.formatting_manager = FormattingManager(self.document)
        self.style_manager = StyleManager(self.document)
        self.ai_integration = AIIntegration()
        
        return self.document
    
    def save_document(self, output_path: str) -> bool:
        """
        保存文档
        
        Args:
            output_path: 输出路径
            
        Returns:
            是否保存成功
        """
        if not self.document:
            raise ValueError("没有打开的文档")
            
        try:
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            self.document.save(output_path)
            return True
        except Exception as e:
            print(f"保存文档失败: {e}")
            return False
    
    def analyze_document(self, use_ai: bool = True) -> Dict[str, Any]:
        """
        分析文档结构
        
        Args:
            use_ai: 是否使用AI分析
            
        Returns:
            分析结果字典
        """
        if not self.document:
            raise ValueError("没有打开的文档")
            
        analysis = {
            "paragraphs": [],
            "tables": [],
            "sections": [],
            "styles": [],
            "statistics": {}
        }
        
        # 分析段落
        for i, para in enumerate(self.document.paragraphs):
            para_info = {
                "index": i,
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
                "alignment": str(para.alignment) if para.alignment else None,
                "runs": []
            }
            
            # 分析runs
            for run in para.runs:
                run_info = {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "font_name": run.font.name,
                    "font_size": run.font.size
                }
                para_info["runs"].append(run_info)
                
            analysis["paragraphs"].append(para_info)
            
        # 分析表格
        for i, table in enumerate(self.document.tables):
            table_info = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cells": []
            }
            
            for row in table.rows:
                for cell in row.cells:
                    cell_info = {
                        "text": cell.text,
                        "row": cell.row.index,
                        "column": cell.column.index
                    }
                    table_info["cells"].append(cell_info)
                    
            analysis["tables"].append(table_info)
            
        # 分析节
        for i, section in enumerate(self.document.sections):
            section_info = {
                "index": i,
                "page_width": section.page_width,
                "page_height": section.page_height,
                "left_margin": section.left_margin,
                "right_margin": section.right_margin,
                "top_margin": section.top_margin,
                "bottom_margin": section.bottom_margin
            }
            analysis["sections"].append(section_info)
            
        # 统计信息
        analysis["statistics"] = {
            "paragraph_count": len(self.document.paragraphs),
            "table_count": len(self.document.tables),
            "section_count": len(self.document.sections),
            "character_count": sum(len(p.text) for p in self.document.paragraphs),
            "word_count": sum(len(p.text.split()) for p in self.document.paragraphs if p.text.strip())
        }
        
        # 使用AI分析（如果启用）
        if use_ai and self.ai_integration:
            ai_analysis = self.ai_integration.analyze_structure(analysis)
            analysis["ai_analysis"] = ai_analysis
            
        return analysis
    
    def extract_text(self, max_paragraphs: Optional[int] = None) -> str:
        """
        提取文档文本内容
        
        Args:
            max_paragraphs: 最大段落数
            
        Returns:
            文档文本内容
        """
        if not self.document:
            raise ValueError("没有打开的文档")
            
        paragraphs = self.document.paragraphs
        if max_paragraphs:
            paragraphs = paragraphs[:max_paragraphs]
            
        text_parts = []
        for para in paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                
        return "\n".join(text_parts)
    
    def get_paragraph_by_index(self, index: int) -> Optional[Any]:
        """
        根据索引获取段落
        
        Args:
            index: 段落索引
            
        Returns:
            段落对象
        """
        if not self.document:
            raise ValueError("没有打开的文档")
            
        if 0 <= index < len(self.document.paragraphs):
            return self.document.paragraphs[index]
            
        return None
    
    def find_paragraphs_by_text(self, text: str, exact_match: bool = False) -> List[Any]:
        """
        根据文本查找段落
        
        Args:
            text: 要查找的文本
            exact_match: 是否精确匹配
            
        Returns:
            段落对象列表
        """
        if not self.document:
            raise ValueError("没有打开的文档")
            
        results = []
        for para in self.document.paragraphs:
            if exact_match:
                if para.text.strip() == text:
                    results.append(para)
            else:
                if text in para.text:
                    results.append(para)
                    
        return results
    
    def replace_text(self, old_text: str, new_text: str, preserve_formatting: bool = True) -> int:
        """
        替换文档中的文本
        
        Args:
            old_text: 原文本
            new_text: 新文本
            preserve_formatting: 是否保留格式
            
        Returns:
            替换的数量
        """
        if not self.document:
            raise ValueError("没有打开的文档")
            
        count = 0
        
        # 替换段落中的文本
        for para in self.document.paragraphs:
            if old_text in para.text:
                if preserve_formatting:
                    # 保留格式的替换
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            count += 1
                else:
                    # 简单替换
                    para.text = para.text.replace(old_text, new_text)
                    count += 1
                    
        # 替换表格中的文本
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        if preserve_formatting:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)
                                        count += 1
                        else:
                            cell.text = cell.text.replace(old_text, new_text)
                            count += 1
                            
        return count
    
    def add_paragraph(self, text: str, style: Optional[str] = None, 
                     alignment: Optional[str] = None) -> Any:
        """
        添加段落
        
        Args:
            text: 段落文本
            style: 样式名称
            alignment: 对齐方式
            
        Returns:
            段落对象
        """
        if not self.document:
            raise ValueError("没有打开的文档")
            
        para = self.document.add_paragraph(text, style)
        
        if alignment:
            alignment_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            if alignment.lower() in alignment_map:
                para.alignment = alignment_map[alignment.lower()]
                
        return para
    
    def add_heading(self, text: str, level: int = 1) -> Any:
        """
        添加标题
        
        Args:
            text: 标题文本
            level: 标题级别
            
        Returns:
            段落对象
        """
        if not self.document:
            raise ValueError("没有打开的文档")
            
        return self.document.add_heading(text, level)
    
    def add_table(self, rows: int, columns: int, data: Optional[List[List[str]]] = None) -> Any:
        """
        添加表格
        
        Args:
            rows: 行数
            columns: 列数
            data: 表格数据
            
        Returns:
            表格对象
        """
        if not self.document:
            raise ValueError("没有打开的文档")
            
        table = self.document.add_table(rows=rows, cols=columns)
        
        if data:
            for i, row_data in enumerate(data):
                if i < rows:
                    for j, cell_text in enumerate(row_data):
                        if j < columns:
                            table.cell(i, j).text = str(cell_text)
                            
        return table
    
    def merge_documents(self, other_doc_path: str, position: str = "end") -> bool:
        """
        合并文档
        
        Args:
            other_doc_path: 另一个文档的路径
            position: 插入位置（start, end）
            
        Returns:
            是否合并成功
        """
        if not self.document:
            raise ValueError("没有打开的文档")
            
        if not os.path.exists(other_doc_path):
            raise FileNotFoundError(f"文档不存在: {other_doc_path}")
            
        try:
            other_doc = Document(other_doc_path)
            
            if position == "start":
                # 在开头插入
                for i, para in enumerate(other_doc.paragraphs):
                    new_para = self.document.paragraphs[0].insert_before(para._element)
            else:
                # 在结尾插入
                for para in other_doc.paragraphs:
                    self.document.add_paragraph(para.text, para.style)
                    
            return True
        except Exception as e:
            print(f"合并文档失败: {e}")
            return False
    
    def get_document_info(self) -> Dict[str, Any]:
        """
        获取文档信息
        
        Returns:
            文档信息字典
        """
        if not self.document:
            raise ValueError("没有打开的文档")
            
        info = {
            "file_path": self.template_path or "新建文档",
            "paragraph_count": len(self.document.paragraphs),
            "table_count": len(self.document.tables),
            "section_count": len(self.document.sections),
            "character_count": sum(len(p.text) for p in self.document.paragraphs),
            "styles": [style.name for style in self.document.styles],
            "core_properties": {}
        }
        
        # 获取核心属性
        if self.document.core_properties:
            props = self.document.core_properties
            info["core_properties"] = {
                "title": props.title,
                "author": props.author,
                "subject": props.subject,
                "keywords": props.keywords,
                "created": str(props.created) if props.created else None,
                "modified": str(props.modified) if props.modified else None
            }
            
        return info