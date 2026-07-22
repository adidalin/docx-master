"""
目录生成模块
提供自动目录生成功能
"""

from typing import Dict, List, Optional, Any
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class TableOfContents:
    """目录生成器"""
    
    def __init__(self, document: Document):
        """
        初始化目录生成器
        
        Args:
            document: Document对象
        """
        self.document = document
    
    def generate(self, title: str = "目录", max_level: int = 3) -> bool:
        """
        生成目录
        
        Args:
            title: 目录标题
            max_level: 最大标题级别
            
        Returns:
            是否成功
        """
        try:
            # 添加目录标题
            toc_title = self.document.add_paragraph(title)
            toc_title.style = self.document.styles['Heading 1']
            
            # 添加目录字段
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run()
            
            # 添加目录字段代码
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._element.append(fldChar1)
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = f' TOC \\o "1-{max_level}" \\h \\z \\u '
            run._element.append(instrText)
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            run._element.append(fldChar2)
            
            # 添加占位文本
            run2 = paragraph.add_run('请更新目录')
            run2.font.color.rgb = None  # 默认颜色
            
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            run2._element.append(fldChar3)
            
            return True
        except Exception as e:
            print(f"生成目录失败: {e}")
            return False
    
    def update(self) -> bool:
        """
        更新目录
        
        Returns:
            是否成功
        """
        try:
            # 查找目录字段
            for paragraph in self.document.paragraphs:
                for run in paragraph.runs:
                    for child in run._element:
                        if child.tag == qn('w:instrText'):
                            if 'TOC' in child.text:
                                # 找到目录，需要更新
                                # 注意：python-docx不能直接更新目录字段
                                # 需要在Word中手动更新或使用其他方法
                                return True
            return False
        except Exception as e:
            print(f"更新目录失败: {e}")
            return False
    
    def add_toc_entry(self, text: str, level: int, page: int = 1) -> bool:
        """
        添加目录条目
        
        Args:
            text: 条目文本
            level: 条目级别
            page: 页码
            
        Returns:
            是否成功
        """
        try:
            # 创建目录条目
            paragraph = self.document.add_paragraph()
            
            # 设置缩进
            paragraph.paragraph_format.left_indent = (level - 1) * 0.5
            
            # 添加文本
            run = paragraph.add_run(text)
            
            # 添加制表符
            run2 = paragraph.add_run('\t')
            
            # 添加页码
            run3 = paragraph.add_run(str(page))
            
            return True
        except Exception as e:
            print(f"添加目录条目失败: {e}")
            return False
    
    def extract_headings(self) -> List[Dict[str, Any]]:
        """
        提取文档标题
        
        Returns:
            标题列表
        """
        headings = []
        
        for i, para in enumerate(self.document.paragraphs):
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading', ''))
                    headings.append({
                        "index": i,
                        "level": level,
                        "text": para.text
                    })
                except ValueError:
                    pass
        
        return headings
    
    def generate_simple_toc(self) -> str:
        """
        生成简单目录文本
        
        Returns:
            目录文本
        """
        headings = self.extract_headings()
        
        toc_lines = []
        for heading in headings:
            indent = "  " * (heading["level"] - 1)
            toc_lines.append(f"{indent}{heading['text']}")
        
        return "\n".join(toc_lines)


def add_toc_to_document(doc_path: str, output_path: str = None, 
                        title: str = "目录", max_level: int = 3) -> bool:
    """
    为文档添加目录
    
    Args:
        doc_path: 输入文档路径
        output_path: 输出文档路径
        title: 目录标题
        max_level: 最大标题级别
        
    Returns:
        是否成功
    """
    try:
        doc = Document(doc_path)
        toc = TableOfContents(doc)
        
        # 在文档开头插入目录
        # 注意：这里需要在文档开头插入，而不是添加到末尾
        toc.generate(title, max_level)
        
        if output_path is None:
            output_path = doc_path
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"添加目录失败: {e}")
        return False


def extract_toc_from_document(doc_path: str) -> str:
    """
    从文档提取目录
    
    Args:
        doc_path: 文档路径
        
    Returns:
        目录文本
    """
    try:
        doc = Document(doc_path)
        toc = TableOfContents(doc)
        return toc.generate_simple_toc()
    except Exception as e:
        print(f"提取目录失败: {e}")
        return ""
