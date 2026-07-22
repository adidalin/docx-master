"""
占位符填充模块
提供占位符查找和填充功能
"""

import re
from typing import Dict, List, Optional, Any
from docx import Document


class PlaceholderFiller:
    """占位符填充器"""
    
    def __init__(self, document: Document):
        """
        初始化占位符填充器
        
        Args:
            document: Document对象
        """
        self.document = document
    
    def find_placeholders(self, pattern: str = r'\{\{([^}]+)\}\}') -> List[Dict[str, Any]]:
        """
        查找占位符
        
        Args:
            pattern: 占位符正则表达式模式
            
        Returns:
            占位符列表
        """
        placeholders = []
        
        for i, para in enumerate(self.document.paragraphs):
            matches = re.finditer(pattern, para.text)
            for match in matches:
                placeholders.append({
                    "index": i,
                    "name": match.group(1),
                    "full_match": match.group(0),
                    "start": match.start(),
                    "end": match.end()
                })
        
        # 检查表格
        for table_idx, table in enumerate(self.document.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        matches = re.finditer(pattern, para.text)
                        for match in matches:
                            placeholders.append({
                                "index": f"table_{table_idx}_{row_idx}_{col_idx}_{para_idx}",
                                "name": match.group(1),
                                "full_match": match.group(0),
                                "start": match.start(),
                                "end": match.end(),
                                "location": "table"
                            })
        
        return placeholders
    
    def fill(self, data: Dict[str, str], pattern: str = r'\{\{([^}]+)\}\}') -> Dict[str, Any]:
        """
        填充占位符
        
        Args:
            data: 占位符数据 {名称: 值}
            pattern: 占位符正则表达式模式
            
        Returns:
            填充结果
        """
        result = {
            "filled": [],
            "not_found": [],
            "errors": []
        }
        
        # 填充段落中的占位符
        for i, para in enumerate(self.document.paragraphs):
            for run in para.runs:
                for name, value in data.items():
                    placeholder = f"{{{{{name}}}}}"
                    if placeholder in run.text:
                        try:
                            run.text = run.text.replace(placeholder, str(value))
                            result["filled"].append({
                                "location": f"paragraph_{i}",
                                "name": name,
                                "value": value
                            })
                        except Exception as e:
                            result["errors"].append({
                                "location": f"paragraph_{i}",
                                "name": name,
                                "error": str(e)
                            })
        
        # 填充表格中的占位符
        for table_idx, table in enumerate(self.document.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        for run in para.runs:
                            for name, value in data.items():
                                placeholder = f"{{{{{name}}}}}"
                                if placeholder in run.text:
                                    try:
                                        run.text = run.text.replace(placeholder, str(value))
                                        result["filled"].append({
                                            "location": f"table_{table_idx}_{row_idx}_{col_idx}_{para_idx}",
                                            "name": name,
                                            "value": value
                                        })
                                    except Exception as e:
                                        result["errors"].append({
                                            "location": f"table_{table_idx}_{row_idx}_{col_idx}_{para_idx}",
                                            "name": name,
                                            "error": str(e)
                                        })
        
        # 检查未找到的占位符
        all_placeholders = self.find_placeholders(pattern)
        for placeholder in all_placeholders:
            if placeholder["name"] not in data:
                result["not_found"].append(placeholder["name"])
        
        return result
    
    def fill_from_json(self, json_path: str, pattern: str = r'\{\{([^}]+)\}\}') -> Dict[str, Any]:
        """
        从JSON文件填充占位符
        
        Args:
            json_path: JSON文件路径
            pattern: 占位符正则表达式模式
            
        Returns:
            填充结果
        """
        try:
            import json
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            return self.fill(data, pattern)
        except Exception as e:
            return {
                "filled": [],
                "not_found": [],
                "errors": [{"error": str(e)}]
            }
    
    def replace_text(self, old_text: str, new_text: str) -> int:
        """
        替换文本
        
        Args:
            old_text: 原文本
            new_text: 新文本
            
        Returns:
            替换数量
        """
        count = 0
        
        # 替换段落中的文本
        for para in self.document.paragraphs:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
        
        # 替换表格中的文本
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                count += 1
        
        return count


def fill_placeholders(doc_path: str, data: Dict[str, str], 
                      output_path: str = None) -> Dict[str, Any]:
    """
    填充文档占位符
    
    Args:
        doc_path: 文档路径
        data: 占位符数据
        output_path: 输出路径
        
    Returns:
        填充结果
    """
    try:
        doc = Document(doc_path)
        filler = PlaceholderFiller(doc)
        
        result = filler.fill(data)
        
        if output_path is None:
            output_path = doc_path
        
        doc.save(output_path)
        return result
    except Exception as e:
        return {
            "filled": [],
            "not_found": [],
            "errors": [{"error": str(e)}]
        }


def fill_from_json(doc_path: str, json_path: str, 
                   output_path: str = None) -> Dict[str, Any]:
    """
    从JSON文件填充文档占位符
    
    Args:
        doc_path: 文档路径
        json_path: JSON文件路径
        output_path: 输出路径
        
    Returns:
        填充结果
    """
    try:
        doc = Document(doc_path)
        filler = PlaceholderFiller(doc)
        
        result = filler.fill_from_json(json_path)
        
        if output_path is None:
            output_path = doc_path
        
        doc.save(output_path)
        return result
    except Exception as e:
        return {
            "filled": [],
            "not_found": [],
            "errors": [{"error": str(e)}]
        }


def replace_text_in_document(doc_path: str, old_text: str, new_text: str,
                             output_path: str = None) -> int:
    """
    替换文档中的文本
    
    Args:
        doc_path: 文档路径
        old_text: 原文本
        new_text: 新文本
        output_path: 输出路径
        
    Returns:
        替换数量
    """
    try:
        doc = Document(doc_path)
        filler = PlaceholderFiller(doc)
        
        count = filler.replace_text(old_text, new_text)
        
        if output_path is None:
            output_path = doc_path
        
        doc.save(output_path)
        return count
    except Exception as e:
        print(f"替换文本失败: {e}")
        return 0
